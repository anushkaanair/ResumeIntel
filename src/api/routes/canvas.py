"""Canvas API routes — real scoring, provenance, dispute, profile sync, export."""

from __future__ import annotations

import io
import re
import uuid
from datetime import datetime, timezone
from typing import Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

from src.db.redis_store import rget, rpush_list, rget_list, rset

router = APIRouter()

# Redis key prefixes
_CANVAS_KEY = "canvas:{resume_id}"
_BULLET_KEY = "bullet:{bullet_id}"
_VERSION_KEY = "versions:{resume_id}"


# ─── Schemas ─────────────────────────────────────────────────────────────────


class BulletScoreRequest(BaseModel):
    text: str = Field(..., min_length=5)
    job_description: str = ""


class BulletSuggestRequest(BaseModel):
    text: str = Field(..., min_length=5)
    job_description: str = ""
    resume_text: str = ""


class BulletAcceptRequest(BaseModel):
    version: str = Field(..., description="'ai' | 'user' | 'original'")
    text: str = ""


class SectionReoptimizeRequest(BaseModel):
    section_id: str
    section_title: str
    section_content: str = ""
    job_description: str = ""
    resume_text: str = ""


class SectionEnhanceRequest(BaseModel):
    section_id: str
    section_title: str
    prompt: str
    section_content: str = ""
    resume_text: str = ""


class ProfileRefreshRequest(BaseModel):
    job_description: str = ""
    last_sync_at: str = ""  # ISO timestamp


class DisputeRequest(BaseModel):
    bullet_text: str
    user_disagreement: str = Field(..., description="Why the user disagrees with the AI score/content")
    job_description: str = ""
    resume_text: str = ""


class ExportRequest(BaseModel):
    resume_id: str
    bullets: list[dict[str, Any]] = Field(default_factory=list)
    sections: list[dict[str, Any]] = Field(default_factory=list)
    filename: str = "optimized_resume"
    format: str = "docx"


# ─── Scoring helpers ──────────────────────────────────────────────────────────

STRONG_VERBS = {
    "led", "designed", "implemented", "optimized", "delivered", "built",
    "architected", "automated", "reduced", "increased", "achieved",
    "launched", "migrated", "scaled", "streamlined", "orchestrated",
    "spearheaded", "developed", "engineered", "deployed", "managed",
}


def _score_bullet(text: str, jd: str = "") -> dict[str, Any]:
    lower = text.lower()
    has_strong_verb = any(v in lower for v in STRONG_VERBS)
    has_metric = bool(re.search(r"\d+", text))
    has_outcome = any(w in lower for w in ["%", "revenue", "users", "latency", "uptime", "reduced", "improved", "saving"])
    impact_score = 0.3 + (0.25 if has_strong_verb else 0) + (0.25 if has_metric else 0) + (0.2 if has_outcome else 0)

    alignment_delta = 0.0
    if jd:
        jd_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", jd.lower()))
        bullet_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", lower))
        stop = {"the", "and", "for", "with", "that", "this", "from"}
        jd_kw = jd_words - stop
        if jd_kw:
            alignment_delta = len(jd_kw & bullet_words) / len(jd_kw)

    return {"impact_score": round(impact_score, 3), "alignment_delta": round(alignment_delta, 3)}


# ─── Routes ──────────────────────────────────────────────────────────────────


@router.get("/canvas/{resume_id}/state")
async def get_canvas_state(resume_id: str) -> dict:
    """Return full canvas state including pipeline output if available."""
    stored = await rget(_CANVAS_KEY.format(resume_id=resume_id))
    if stored:
        return {"status": "ok", "data": stored}

    # Check if a job result exists under this id (optimize stores under job_id)
    job_data = await rget(f"job:{resume_id}")
    if job_data and job_data.get("status") == "completed":
        await rset(_CANVAS_KEY.format(resume_id=resume_id), job_data)
        return {"status": "ok", "data": job_data}

    return {
        "status": "ok",
        "data": {
            "resume_id": resume_id,
            "status": "pending",
            "message": "Pipeline output not yet available. Connect WebSocket for live updates.",
        },
    }


@router.post("/canvas/bullet/{bullet_id}/score")
async def score_bullet(bullet_id: str, request: BulletScoreRequest) -> dict:
    """Real impact scoring using heuristics + keyword alignment."""
    scores = _score_bullet(request.text, request.job_description)
    return {"status": "ok", "data": scores}


@router.post("/canvas/bullet/{bullet_id}/suggest")
async def suggest_bullet(bullet_id: str, request: BulletSuggestRequest) -> dict:
    """Generate an AI improvement suggestion using the QualityAgent LLM."""
    from src.llm.client import LLMClient
    from src.rag.embedder import Embedder
    from src.rag.retriever import Retriever
    from src.rag.vector_store import VectorStore

    llm = LLMClient()
    embedder = Embedder()
    vs = VectorStore()
    retriever = Retriever(embedder, vs, user_id=bullet_id)

    if request.resume_text:
        segments = [
            {"content": s.strip(), "segment_id": f"s{i}", "section": "resume"}
            for i, s in enumerate(request.resume_text.split("\n"))
            if s.strip() and len(s.strip()) > 15
        ]
        await retriever.index_segments(segments)

    context = await retriever.retrieve(request.text, top_k=3)
    context_text = "\n".join(f"- {seg.content}" for seg in context)

    prompt = f"""You are a senior resume editor. Rewrite this bullet to be more impactful.

CURRENT BULLET:
{request.text}

JOB DESCRIPTION CONTEXT:
{request.job_description[:500] if request.job_description else 'Not provided'}

SOURCE DATA (use only these facts — no fabrication):
{context_text}

Rules:
- Start with a strong past-tense action verb
- Include at least one metric if present in source data
- Keep under 25 words
- Return ONLY the rewritten bullet, no explanation"""

    suggestion = await llm.generate(prompt)
    suggestion = suggestion.strip().lstrip("-•* ")

    chunk_ids = [seg.segment_id for seg in context]
    return {
        "status": "ok",
        "data": {
            "suggestion": suggestion,
            "rationale": "Rewritten with strong action verb, metrics from source data, and outcome framing.",
            "provenance": {
                "agent_name": "QualityAgent",
                "retrieved_chunk_ids": chunk_ids,
                "confidence": _score_bullet(suggestion)["impact_score"],
            },
        },
    }


@router.post("/canvas/bullet/{bullet_id}/accept")
async def accept_bullet(bullet_id: str, request: BulletAcceptRequest) -> dict:
    """Record human decision state for a bullet (accept/reject/user_modified)."""
    status_map = {"ai": "accepted", "user": "user_modified", "original": "rejected"}
    status = status_map.get(request.version, "pending")

    await rset(_BULLET_KEY.format(bullet_id=bullet_id), {
        "id": bullet_id,
        "text": request.text,
        "status": status,
        "decided_at": datetime.now(timezone.utc).isoformat(),
    })

    return {
        "status": "ok",
        "data": {"id": bullet_id, "status": status},
    }


@router.post("/canvas/dispute/{bullet_id}")
async def dispute_bullet(bullet_id: str, request: DisputeRequest) -> dict:
    """Human-initiated re-evaluation loop (patent claim 4).

    Re-evaluates the bullet with user disagreement incorporated as context.
    """
    from src.llm.client import LLMClient
    from src.rag.embedder import Embedder
    from src.rag.retriever import Retriever
    from src.rag.vector_store import VectorStore

    llm = LLMClient()
    embedder = Embedder()
    vs = VectorStore()
    retriever = Retriever(embedder, vs, user_id=f"dispute_{bullet_id}")

    if request.resume_text:
        segments = [
            {"content": s.strip(), "segment_id": f"s{i}", "section": "resume"}
            for i, s in enumerate(request.resume_text.split("\n"))
            if s.strip() and len(s.strip()) > 15
        ]
        await retriever.index_segments(segments)

    context = await retriever.retrieve(request.bullet_text, top_k=5)
    context_text = "\n".join(f"- {seg.content}" for seg in context)

    prompt = f"""You are a senior resume quality reviewer re-evaluating a bullet after user dispute.

BULLET:
{request.bullet_text}

USER DISAGREEMENT:
{request.user_disagreement}

SOURCE DATA:
{context_text}

JOB DESCRIPTION:
{request.job_description[:500] if request.job_description else 'Not provided'}

Re-evaluate and provide:
1. A revised quality score (0.0-1.0)
2. A revised or confirmed version of the bullet
3. A rationale addressing the user's disagreement specifically

Format your response as:
SCORE: <float>
BULLET: <revised bullet text>
RATIONALE: <one paragraph explaining your evaluation>"""

    raw = await llm.generate(prompt)

    score = 0.7
    bullet_out = request.bullet_text
    rationale = raw

    for line in raw.splitlines():
        if line.startswith("SCORE:"):
            try:
                score = float(line.split(":", 1)[1].strip())
            except ValueError:
                pass
        elif line.startswith("BULLET:"):
            bullet_out = line.split(":", 1)[1].strip()
        elif line.startswith("RATIONALE:"):
            rationale = line.split(":", 1)[1].strip()

    chunk_ids = [seg.segment_id for seg in context]
    return {
        "status": "ok",
        "data": {
            "bullet_id": bullet_id,
            "revised_bullet": bullet_out,
            "revised_score": round(score, 3),
            "provenance": {
                "agent_name": "QualityAgent",
                "input_summary": request.bullet_text[:200],
                "retrieved_chunk_ids": chunk_ids,
                "decision_rationale": rationale,
                "confidence": round(score, 3),
            },
        },
    }


@router.post("/canvas/section/{section_id}/reoptimize")
async def reoptimize_section(section_id: str, request: SectionReoptimizeRequest) -> dict:
    """Re-optimize a section using QualityAgent via real LLM call."""
    from src.llm.client import LLMClient

    if not request.section_content:
        return {"status": "ok", "data": {"message": "No section content provided", "section_id": section_id}}

    llm = LLMClient()
    prompt = f"""You are a senior resume optimizer. Rewrite this resume section to be stronger.

SECTION: {request.section_title}
CONTENT:
{request.section_content}

JOB DESCRIPTION:
{request.job_description[:600] if request.job_description else 'Not provided'}

Rules:
- Replace weak verbs with strong action verbs
- Add metrics where present in content
- Keep all factual details intact — no fabrication
- Return ONLY the rewritten section content in the same format"""

    rewritten = await llm.generate(prompt)
    return {
        "status": "ok",
        "data": {
            "section_id": section_id,
            "rewritten_content": rewritten.strip(),
            "message": f"{request.section_title} reoptimized by QualityAgent",
        },
    }


@router.post("/canvas/section/{section_id}/enhance")
async def enhance_section(section_id: str, request: SectionEnhanceRequest) -> dict:
    """Apply a user prompt to enhance a section."""
    from src.llm.client import LLMClient

    llm = LLMClient()
    prompt = f"""You are a resume editor. Apply this specific instruction to the section below.

INSTRUCTION: {request.prompt}

SECTION: {request.section_title}
CONTENT:
{request.section_content or 'Not provided'}

Return ONLY the enhanced section content. No explanation."""

    enhanced = await llm.generate(prompt)
    return {
        "status": "ok",
        "data": {
            "section_id": section_id,
            "enhanced_content": enhanced.strip(),
            "message": f"{request.section_title} enhanced.",
        },
    }


# ─── Profile Sync Engine (patent claim 1d) ────────────────────────────────────


@router.post("/canvas/profile/linkedin/refresh")
async def refresh_linkedin(request: ProfileRefreshRequest) -> dict:
    """Staleness-aware LinkedIn profile sync with JD-semantic filtering (patent claim 1d)."""
    from src.config.settings import settings
    from src.rag.embedder import Embedder

    embedder = Embedder()
    staleness_score = _compute_staleness(request.last_sync_at)

    raw_deltas = await _fetch_linkedin_deltas(settings.linkedin_client_id, settings.linkedin_client_secret)
    filtered = _filter_by_jd_relevance(embedder, raw_deltas, request.job_description, threshold=0.5)

    return {
        "status": "ok",
        "data": {
            "platform": "linkedin",
            "staleness_score": round(staleness_score, 3),
            "has_changes": len(filtered) > 0,
            "items": filtered,
            "last_sync_at": datetime.now(timezone.utc).isoformat(),
        },
    }


@router.post("/canvas/profile/github/refresh")
async def refresh_github(request: ProfileRefreshRequest) -> dict:
    """Staleness-aware GitHub profile sync with JD-semantic filtering (patent claim 1d)."""
    from src.config.settings import settings
    from src.rag.embedder import Embedder

    embedder = Embedder()
    staleness_score = _compute_staleness(request.last_sync_at)

    raw_deltas = await _fetch_github_deltas(settings.github_access_token)
    filtered = _filter_by_jd_relevance(embedder, raw_deltas, request.job_description, threshold=0.5)

    return {
        "status": "ok",
        "data": {
            "platform": "github",
            "staleness_score": round(staleness_score, 3),
            "has_changes": len(filtered) > 0,
            "items": filtered,
            "last_sync_at": datetime.now(timezone.utc).isoformat(),
        },
    }


async def _fetch_linkedin_deltas(client_id: str, client_secret: str) -> list[dict]:
    """Fetch career updates from LinkedIn API.

    Requires LINKEDIN_CLIENT_ID and LINKEDIN_CLIENT_SECRET in .env.
    Returns empty list if credentials are not configured.
    """
    if not client_id or not client_secret:
        return []

    import httpx
    # LinkedIn API v2 — positions and certifications endpoint
    # NOTE: requires a valid OAuth access token per user session (3-legged OAuth).
    # The token should be passed from the frontend after the user authorises the app.
    # For now returns empty; wire frontend OAuth flow to pass token in request headers.
    return []


async def _fetch_github_deltas(access_token: str) -> list[dict]:
    """Fetch recent repos and activity from GitHub API.

    Requires GITHUB_ACCESS_TOKEN in .env (fine-grained personal access token,
    read:user + read:repo scopes).
    Returns empty list if token is not configured.
    """
    if not access_token:
        return []

    import httpx
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }

    deltas: list[dict] = []
    async with httpx.AsyncClient(timeout=10) as client:
        # Fetch recent repos
        resp = await client.get("https://api.github.com/user/repos?sort=updated&per_page=10", headers=headers)
        if resp.status_code == 200:
            for repo in resp.json():
                if not repo.get("fork"):
                    lang = repo.get("language") or ""
                    desc = repo.get("description") or ""
                    text = f"Repo: {repo['name']} ({lang}){' — ' + desc if desc else ''}"
                    deltas.append({"type": "repo", "text": text, "platform": "github", "url": repo.get("html_url", "")})

        # Fetch user languages from recent repos
        resp2 = await client.get("https://api.github.com/user", headers=headers)
        if resp2.status_code == 200:
            user = resp2.json()
            bio = user.get("bio") or ""
            if bio:
                deltas.append({"type": "bio", "text": bio, "platform": "github"})

    return deltas


def _compute_staleness(last_sync_at: str) -> float:
    """min(days_since_sync / 30, 1.0) per patent claim 5."""
    if not last_sync_at:
        return 1.0
    try:
        last = datetime.fromisoformat(last_sync_at.replace("Z", "+00:00"))
        days = (datetime.now(timezone.utc) - last).days
        return min(days / 30.0, 1.0)
    except Exception:
        return 1.0


def _filter_by_jd_relevance(
    embedder, deltas: list[dict], jd: str, threshold: float = 0.5
) -> list[dict]:
    """Filter delta items by SBERT cosine similarity to active JD."""
    if not jd or not deltas:
        return deltas  # no filtering without JD

    import numpy as np

    jd_vec = embedder.encode_single(jd)
    result = []
    for item in deltas:
        item_vec = embedder.encode_single(item["text"])
        score = float(np.dot(item_vec, jd_vec))
        if score >= threshold:
            result.append({**item, "relevance_score": round(score, 3)})
    return result


# ─── Decision-state-conditional export (patent claim 5, 13) ──────────────────


@router.post("/canvas/export")
async def export_canvas(request: ExportRequest) -> Response:
    """Assemble export from ONLY accepted/user_modified bullets (patent claim 13)."""
    # Build content from decision-state-gated bullets
    content_lines: list[str] = []

    for section in request.sections:
        content_lines.append(f"## {section.get('title', 'Section')}\n")
        entries = section.get("content", {}).get("entries", [])
        summary_text = section.get("content", {}).get("text", "")

        if summary_text:
            content_lines.append(summary_text + "\n")

        for entry in entries:
            if entry.get("company") or entry.get("title"):
                content_lines.append(
                    f"**{entry.get('title', '')}** — {entry.get('company', '')} "
                    f"({entry.get('startDate', '')} – {entry.get('endDate', '')})\n"
                )
            for bullet in entry.get("bullets", []):
                # Decision-state gate: only accepted or user_modified
                status = bullet.get("status", "pending")
                if status not in ("accepted", "user_modified"):
                    continue
                text = bullet.get("currentText") or bullet.get("originalText", "")
                if text:
                    content_lines.append(f"- {text}\n")

    full_content = "\n".join(content_lines)

    if request.format == "pdf":
        return await _build_pdf(full_content, request.filename)
    else:
        return await _build_docx(full_content, request.filename)


async def _build_pdf(content: str, filename: str) -> Response:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        story = []

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.1 * inch))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], styles["Heading2"]))
            elif stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], styles["Heading1"]))
            elif stripped.startswith(("- ", "* ")):
                story.append(Paragraph(f"• {stripped[2:]}", styles["Normal"]))
            else:
                story.append(Paragraph(stripped, styles["Normal"]))

        doc.build(story)
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="reportlab required: pip install reportlab")


# ─── Version history + diff (patent para 21) ─────────────────────────────────


class VersionSnapshotRequest(BaseModel):
    resume_id: str
    content: str
    change_source: str = "user_edit"  # 'accept' | 'reject' | 'user_edit'


class VersionRevertRequest(BaseModel):
    version_id: str


@router.post("/canvas/{resume_id}/version/snapshot")
async def create_snapshot(resume_id: str, request: VersionSnapshotRequest) -> dict:
    """Create a canvas version snapshot (auto-called on every accept/reject/edit)."""
    existing = await rget_list(_VERSION_KEY.format(resume_id=resume_id))
    version_num = len(existing) + 1
    version_id = f"v{version_num}"
    await rpush_list(_VERSION_KEY.format(resume_id=resume_id), {
        "id": version_id,
        "version_num": version_num,
        "content": request.content,
        "change_source": request.change_source,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })
    return {"status": "ok", "data": {"version_id": version_id, "version_num": version_num}}


@router.get("/canvas/{resume_id}/version/history")
async def get_version_history(resume_id: str) -> dict:
    snapshots = await rget_list(_VERSION_KEY.format(resume_id=resume_id))
    return {
        "status": "ok",
        "data": {
            "resume_id": resume_id,
            "versions": [
                {k: v for k, v in s.items() if k != "content"}
                for s in snapshots
            ],
        },
    }


@router.get("/canvas/{resume_id}/version/{version_id}/diff")
async def get_version_diff(resume_id: str, version_id: str) -> dict:
    """Return structured diff between a version and the next one, + AI summary."""
    import difflib
    from src.llm.client import LLMClient

    snapshots = await rget_list(_VERSION_KEY.format(resume_id=resume_id))
    target = next((s for s in snapshots if s["id"] == version_id), None)
    if not target:
        raise HTTPException(status_code=404, detail=f"Version {version_id} not found")

    idx = snapshots.index(target)
    next_snap = snapshots[idx + 1] if idx + 1 < len(snapshots) else None

    if not next_snap:
        return {"status": "ok", "data": {"diff": [], "summary": "This is the latest version."}}

    a_lines = target["content"].splitlines(keepends=True)
    b_lines = next_snap["content"].splitlines(keepends=True)

    diff_ops: list[dict] = []
    for op, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a_lines, b_lines).get_opcodes():
        if op == "equal":
            for line in a_lines[i1:i2]:
                diff_ops.append({"type": "same", "text": line.rstrip()})
        elif op in ("replace", "delete"):
            for line in a_lines[i1:i2]:
                diff_ops.append({"type": "remove", "text": line.rstrip()})
            if op == "replace":
                for line in b_lines[j1:j2]:
                    diff_ops.append({"type": "add", "text": line.rstrip()})
        elif op == "insert":
            for line in b_lines[j1:j2]:
                diff_ops.append({"type": "add", "text": line.rstrip()})

    # AI-generated natural language summary
    added = [d["text"] for d in diff_ops if d["type"] == "add"][:5]
    removed = [d["text"] for d in diff_ops if d["type"] == "remove"][:5]

    llm = LLMClient()
    summary_prompt = f"""Summarize these resume changes in one concise sentence (max 20 words).

ADDED:
{chr(10).join(added) if added else 'Nothing added'}

REMOVED:
{chr(10).join(removed) if removed else 'Nothing removed'}

Return ONLY the summary sentence."""
    summary = (await llm.generate(summary_prompt)).strip()

    return {
        "status": "ok",
        "data": {
            "from_version": version_id,
            "to_version": next_snap["id"],
            "diff": diff_ops,
            "summary": summary,
        },
    }


@router.post("/canvas/{resume_id}/version/revert")
async def revert_to_version(resume_id: str, request: VersionRevertRequest) -> dict:
    """Revert canvas to a historical snapshot."""
    snapshots = await rget_list(_VERSION_KEY.format(resume_id=resume_id))
    target = next((s for s in snapshots if s["id"] == request.version_id), None)
    if not target:
        raise HTTPException(status_code=404, detail=f"Version {request.version_id} not found")
    return {
        "status": "ok",
        "data": {
            "version_id": request.version_id,
            "content": target["content"],
            "timestamp": target["timestamp"],
        },
    }


class QuickATSRequest(BaseModel):
    resume_text: str
    job_description: str = ""


@router.post("/canvas/ats/quick-score")
async def quick_ats_score(request: QuickATSRequest) -> dict:
    """Lightweight ATS score calculation for live recalculation on bullet changes."""
    text = request.resume_text.lower()
    jd = request.job_description.lower()

    jd_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", jd)) if jd else set()
    resume_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", text))
    stop = {"the", "and", "for", "with", "that", "this", "from", "have", "been", "will", "are"}
    jd_kw = jd_words - stop

    keyword_score = (len(jd_kw & resume_words) / len(jd_kw)) if jd_kw else 0.7

    # Format heuristics
    bullet_count = len(re.findall(r"^[\-•]", request.resume_text, re.MULTILINE))
    has_metrics = bool(re.search(r"\d+[%+]?", request.resume_text))
    has_action_verbs = any(v in text for v in STRONG_VERBS)
    has_sections = len(re.findall(r"^#{1,2}\s", request.resume_text, re.MULTILINE)) >= 3

    format_score = (
        0.25 * min(bullet_count / 8, 1.0)
        + 0.25 * (1.0 if has_metrics else 0.0)
        + 0.25 * (1.0 if has_action_verbs else 0.0)
        + 0.25 * (1.0 if has_sections else 0.0)
    )

    ats_score = round((keyword_score * 0.65 + format_score * 0.35) * 100, 1)

    return {
        "status": "ok",
        "data": {
            "ats_score": ats_score,
            "keyword_match_pct": round(keyword_score * 100, 1),
            "format_score": round(format_score * 100, 1),
            "matched_keywords": list(jd_kw & resume_words)[:20],
            "missing_keywords": list(jd_kw - resume_words)[:20],
        },
    }


async def _build_docx(content: str, filename: str) -> Response:
    try:
        from docx import Document

        doc = Document()
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx required: pip install python-docx")
