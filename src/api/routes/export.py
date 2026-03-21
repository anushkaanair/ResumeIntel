"""Export routes — resume export as PDF or DOCX."""
from __future__ import annotations

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

router = APIRouter()


class ExportRequest(BaseModel):
    content: str = Field(
        ..., min_length=10, description="Optimized resume markdown content"
    )
    filename: str = Field(
        default="optimized_resume", description="Output filename without extension"
    )


@router.post("/export/pdf")
async def export_pdf(request: ExportRequest) -> Response:
    """Export optimized resume as PDF."""
    try:
        import io

        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )

        styles = getSampleStyleSheet()
        story = []

        for line in request.content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.1 * inch))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], styles["Heading2"]))
            elif stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], styles["Heading1"]))
            elif stripped.startswith(("- ", "* ", "• ")):
                bullet_text = f"• {stripped[2:]}"
                story.append(Paragraph(bullet_text, styles["Normal"]))
            else:
                story.append(Paragraph(stripped, styles["Normal"]))

        doc.build(story)
        pdf_bytes = buffer.getvalue()

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{request.filename}.pdf"'
            },
        )
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="PDF export requires reportlab: pip install reportlab",
        )


@router.post("/export/docx")
async def export_docx(request: ExportRequest) -> Response:
    """Export optimized resume as DOCX."""
    try:
        import io

        from docx import Document

        doc = Document()

        for line in request.content.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ", "• ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{request.filename}.docx"'
            },
        )
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="DOCX export requires python-docx: pip install python-docx",
        )
